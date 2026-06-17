import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_left_border(cell, color_hex, size_pt):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(int(size_pt * 8)))  # sz is in eighths of a point
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), color_hex)
    tcBorders.append(left)
    
    for border_name in ['top', 'bottom', 'right']:
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'none')
        tcBorders.append(b)
        
    tcPr.append(tcBorders)

def add_formatted_text(paragraph, text):
    # Regex to split on bold (**, ***), italic (*, _), inline code (`) and urls
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`|https?://\S+)')
    parts = pattern.split(text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('***') and part.endswith('***'):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(199, 37, 78) # pinkish code color
        elif part.startswith('http'):
            run = paragraph.add_run(part)
            run.font.color.rgb = RGBColor(0, 102, 204)
            run.underline = True
        else:
            paragraph.add_run(part)

def convert_md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found.")
        return False
        
    doc = Document()
    
    # Configure page margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles config
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(33, 37, 41) # Charcoal
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    code_content = []
    in_table = False
    table_rows = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Handle Code Blocks
        if stripped.startswith('```'):
            if in_code_block:
                # End of code block
                in_code_block = False
                # Write code block as a single paragraph with gray background
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.right_indent = Inches(0.4)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                
                # We can place it in a single cell table for shading
                tbl = doc.add_table(rows=1, cols=1)
                tbl.autofit = False
                cell = tbl.cell(0, 0)
                set_cell_background(cell, "F8F9FA")
                set_cell_left_border(cell, "6C757D", 3)
                
                cp = cell.paragraphs[0]
                cp.paragraph_format.space_before = Pt(4)
                cp.paragraph_format.space_after = Pt(4)
                
                code_text = "\n".join(code_content)
                run = cp.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9.0)
                run.font.color.rgb = RGBColor(40, 44, 52)
                
                code_content = []
            else:
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_content.append(line.rstrip('\n'))
            i += 1
            continue
            
        # Handle Tables
        if stripped.startswith('|'):
            in_table = True
            # Check for table separator row |---|---|
            if re.match(r'^\s*\|(?:\s*:?-+:?\s*\|)+\s*$', stripped):
                i += 1
                continue
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            # Table ended
            if table_rows:
                num_cols = len(table_rows[0])
                num_rows = len(table_rows)
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Table Grid'
                table.autofit = True
                
                for r_idx, row_cells in enumerate(table_rows):
                    row = table.rows[r_idx]
                    for c_idx, cell_value in enumerate(row_cells):
                        if c_idx < len(row.cells):
                            cell = row.cells[c_idx]
                            cell.text = ""
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_before = Pt(3)
                            p.paragraph_format.space_after = Pt(3)
                            add_formatted_text(p, cell_value)
                            
                            # Styling Header Row
                            if r_idx == 0:
                                set_cell_background(cell, "2E5E3B") # EMA forest green
                                for run in p.runs:
                                    run.bold = True
                                    run.font.color.rgb = RGBColor(255, 255, 255)
                            else:
                                if r_idx % 2 == 0:
                                    set_cell_background(cell, "F4F9F5") # Light mint shading
                # Add spacing after table
                p_space = doc.add_paragraph()
                p_space.paragraph_format.space_after = Pt(6)
                table_rows = []
            in_table = False
            
        # Handle Empty Lines
        if not stripped:
            i += 1
            continue
            
        # Handle Headers
        header_match = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if header_match:
            level = len(header_match.group(1))
            title_text = header_match.group(2)
            
            p = doc.add_heading(level=level)
            run = p.add_run(title_text)
            run.font.name = 'Calibri Light'
            
            if level == 1:
                run.font.size = Pt(20)
                run.bold = True
                run.font.color.rgb = RGBColor(46, 94, 59) # Forest Green
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.keep_with_next = True
            elif level == 2:
                run.font.size = Pt(16)
                run.bold = True
                run.font.color.rgb = RGBColor(46, 94, 59)
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.keep_with_next = True
            elif level == 3:
                run.font.size = Pt(13)
                run.bold = True
                run.font.color.rgb = RGBColor(60, 110, 75)
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.keep_with_next = True
            else:
                run.font.size = Pt(11)
                run.bold = True
                run.font.color.rgb = RGBColor(33, 37, 41)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.keep_with_next = True
                
            i += 1
            continue
            
        # Handle Images
        img_match = re.match(r'^!\[(.*?)\]\((.*?)\)', stripped)
        if img_match:
            alt_text = img_match.group(1)
            img_path = img_match.group(2)
            
            if os.path.exists(img_path):
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(4)
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(5.5))
                    
                    caption_p = doc.add_paragraph()
                    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_p.paragraph_format.space_after = Pt(12)
                    c_run = caption_p.add_run(f"Figure: {alt_text}")
                    c_run.italic = True
                    c_run.font.size = Pt(9.5)
                    c_run.font.color.rgb = RGBColor(108, 117, 125)
                except Exception as e:
                    doc.add_paragraph(f"[Failed to embed image '{img_path}': {str(e)}]")
            else:
                doc.add_paragraph(f"[Image file not found: {img_path}]")
            i += 1
            continue
            
        # Handle Blockquotes and Alert Panels
        if stripped.startswith('>'):
            # Check for Alert Type
            alert_type = "quote"
            alert_text = stripped[1:].strip()
            
            if alert_text.startswith('[!NOTE]'):
                alert_type = "note"
                alert_text = alert_text[7:].strip()
            elif alert_text.startswith('[!WARNING]'):
                alert_type = "warning"
                alert_text = alert_text[10:].strip()
            elif alert_text.startswith('[!IMPORTANT]'):
                alert_type = "important"
                alert_text = alert_text[12:].strip()
                
            # Check if next lines are also part of blockquote
            bq_lines = [alert_text]
            next_idx = i + 1
            while next_idx < len(lines) and lines[next_idx].strip().startswith('>'):
                bq_lines.append(lines[next_idx].strip()[1:].strip())
                next_idx += 1
            i = next_idx
            
            full_bq_text = " ".join(bq_lines)
            
            # Create callout table
            tbl = doc.add_table(rows=1, cols=1)
            tbl.autofit = False
            cell = tbl.cell(0, 0)
            
            if alert_type == "note":
                set_cell_background(cell, "F0F7FF") # Soft blue
                set_cell_left_border(cell, "0066CC", 3)
                p = cell.paragraphs[0]
                run_badge = p.add_run("ℹ️ NOTE: ")
                run_badge.bold = True
                run_badge.font.color.rgb = RGBColor(0, 102, 204)
            elif alert_type in ["warning", "important"]:
                set_cell_background(cell, "FFF3CD") # Soft amber
                set_cell_left_border(cell, "D97706", 3)
                p = cell.paragraphs[0]
                run_badge = p.add_run("⚠️ WARNING: ")
                run_badge.bold = True
                run_badge.font.color.rgb = RGBColor(217, 119, 6)
            else:
                set_cell_background(cell, "F8F9FA") # Light gray
                set_cell_left_border(cell, "6C757D", 3)
                p = cell.paragraphs[0]
                
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            add_formatted_text(p, full_bq_text)
            continue
            
        # Handle List Items (Bulleted)
        bullet_match = re.match(r'^([\*\-\+])\s+(.*)$', stripped)
        if bullet_match:
            item_text = bullet_match.group(2)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            add_formatted_text(p, item_text)
            i += 1
            continue
            
        # Handle List Items (Numbered)
        num_match = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        if num_match:
            item_text = num_match.group(2)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
            add_formatted_text(p, item_text)
            i += 1
            continue
            
        # Handle regular paragraphs
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        add_formatted_text(p, stripped)
        i += 1
        
    doc.save(docx_path)
    print(f"Successfully converted MD to DOCX: {docx_path}")
    return True

if __name__ == '__main__':
    base_name = 'SYSTEM_GUIDE.docx'
    success = False
    
    try:
        convert_md_to_docx('SYSTEM_GUIDE.md', base_name)
        success = True
    except PermissionError:
        print(f"Warning: {base_name} is open/locked.")
        
    if not success:
        try:
            convert_md_to_docx('SYSTEM_GUIDE.md', 'SYSTEM_GUIDE_updated.docx')
            success = True
        except PermissionError:
            print("Warning: SYSTEM_GUIDE_updated.docx is also open/locked.")
            
    if not success:
        n = 3
        while not success and n <= 100:
            target_name = f'SYSTEM_GUIDE_v{n}.docx'
            try:
                convert_md_to_docx('SYSTEM_GUIDE.md', target_name)
                success = True
            except PermissionError:
                print(f"Warning: {target_name} is also open/locked.")
                n += 1



